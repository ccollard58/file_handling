import os
import subprocess
import tempfile
import pytesseract
from PIL import Image, ImageEnhance
from pypdf import PdfReader
import logging
from docx import Document
import pandas as pd
import re
import nltk
from nltk import download

# NLTK setup for text quality checking
try:
    nltk.data.find('corpora/words')
except LookupError:
    logging.info("NLTK 'words' corpus not found. Downloading...")
    download('words', quiet=True)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    logging.info("NLTK 'punkt' tokenizer not found. Downloading...")
    download('punkt', quiet=True)

from nltk.corpus import words as nltk_words
from nltk.tokenize import word_tokenize

# Import pdf2image with fallback handling
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    logging.warning("pdf2image not available. PDF OCR will be disabled.")

class DocumentProcessor:
    def __init__(self, llm_analyzer=None):
        # Configure pytesseract path for Windows
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        
        # Setup Poppler path for PDF OCR if provided via environment
        poppler_env = os.getenv('POPPLER_PATH')
        self.poppler_path = poppler_env if poppler_env else None
        
        # Setup LibreOffice (soffice) path for .doc conversion fallback
        self.soffice_path = os.getenv('LIBREOFFICE_PATH')
        if not self.soffice_path:
            # Try common Windows install paths
            possible_paths = [
                r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
                r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
            ]
            for p in possible_paths:
                if os.path.isfile(p):
                    self.soffice_path = p
                    break
        
        # Check if PDF OCR is available
        self.pdf_ocr_available = PDF2IMAGE_AVAILABLE
        
        # Store reference to LLM analyzer for vision model access
        self.llm_analyzer = llm_analyzer
    
    def extract_text(self, file_path):
        """Extract text from PDF or image file."""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.tiff', '.tif']:
                return self._extract_text_from_image(file_path)
            elif file_ext == '.docx':
                return self._extract_text_from_docx(file_path)
            elif file_ext == '.doc':
                return self._extract_text_from_doc(file_path)
            elif file_ext == '.xlsx':
                return self._extract_text_from_xlsx(file_path)
            else:
                logging.warning(f"Unsupported file format: {file_ext}")
                return ""
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_pdf(self, file_path):
        """Extract text from PDF file."""
        text = ""
        try:
            pdf = PdfReader(file_path)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
            
            # If no text was extracted, try OCR
            if not text.strip():
                logging.info(f"No text found in PDF, attempting OCR: {file_path}")
                ocr_text = self._perform_ocr_on_pdf(file_path)
                if ocr_text:
                    text = ocr_text
                else:
                    logging.warning(f"OCR completed for PDF {file_path} but no readable text was found. Attempting vision LLM analysis.")
                    vision_text = self._analyze_pdf_with_vision_llm(file_path)
                    if vision_text:
                        text = vision_text
            # Check if the extracted text contains a reasonable number of real words
            elif not self._has_real_words(text):
                logging.info(f"Extracted PDF text appears to be garbage, attempting OCR: {file_path}")
                ocr_text = self._perform_ocr_on_pdf(file_path)
                if ocr_text and self._has_real_words(ocr_text):
                    text = ocr_text
                else:
                    logging.warning(f"OCR completed for PDF {file_path} but text quality is poor. Attempting vision LLM analysis.")
                    vision_text = self._analyze_pdf_with_vision_llm(file_path)
                    if vision_text:
                        text = vision_text
                    elif ocr_text:
                        text = ocr_text  # Use OCR text as fallback even if quality is poor
                
            return text
        except Exception as e:
            logging.error(f"Error processing PDF {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_image(self, file_path):
        """Extract text from image file using OCR."""
        try:
            image = Image.open(file_path)

            # Handle specific image formats
            if image.format == 'GIF':
                image = image.convert('RGB')
            
            # Special handling for TIFF images, which can sometimes cause OCR issues
            if file_path.lower().endswith(('.tif', '.tiff')):
                # Some TIFFs are multi-page; iterate frames when present
                total_frames = getattr(image, 'n_frames', 1)
                text_chunks = []
                
                for frame_index in range(total_frames):
                    try:
                        if total_frames > 1:
                            image.seek(frame_index)
                        frame = image
                        
                        # Convert to a suitable mode for OCR
                        if frame.mode not in ('RGB', 'L'):
                            frame = frame.convert('RGB')
                        
                        # Log the TIFF processing
                        logging.info(f"Processing TIFF image: {file_path}, frame {frame_index+1}/{total_frames}, mode: {frame.mode}, size: {frame.size}")
                        
                        # Some TIFF files might benefit from image enhancement for OCR
                        try:
                            enhancer = ImageEnhance.Contrast(frame)
                            frame = enhancer.enhance(1.5)  # Boost contrast by 50%
                        except Exception as e:
                            logging.warning(f"Could not enhance TIFF image frame {frame_index+1}: {str(e)}")
                        
                        frame_text = pytesseract.image_to_string(frame)
                        if frame_text and frame_text.strip():
                            text_chunks.append(frame_text)
                    except EOFError:
                        # Safety for malformed multi-page TIFFs
                        break
                    except Exception as e:
                        logging.warning(f"Error processing TIFF frame {frame_index+1}: {e}")
                        continue
                
                text = "\n".join(text_chunks)
            else:
                # Non-TIFF images processed directly
                text = pytesseract.image_to_string(image)
            
            # For image files, add logging to help diagnose OCR issues
            if text.strip():
                text_preview = text.strip()[:50] + ('...' if len(text.strip()) > 50 else '')
                logging.info(f"OCR extracted {len(text.strip())} chars from '{file_path}': {text_preview}")
                
                # Check if OCR text quality is good
                if not self._has_real_words(text):
                    logging.info(f"OCR text quality appears poor for {file_path}, attempting vision LLM analysis")
                    vision_text = self._analyze_image_with_vision_llm(file_path)
                    if vision_text:
                        text = vision_text
            else:
                logging.info(f"OCR extracted no text from '{file_path}', attempting vision LLM analysis")
                vision_text = self._analyze_image_with_vision_llm(file_path)
                if vision_text:
                    text = vision_text

            return text
        except Exception as e:
            logging.error(f"Error processing image {file_path}: {str(e)}")
            return ""
    
    def _perform_ocr_on_pdf(self, file_path):
        """Perform OCR on a PDF file by converting its pages to images."""
        if not self.pdf_ocr_available:
            logging.warning("PDF OCR not available. Please install pdf2image and Poppler.")
            logging.warning("Install with: pip install pdf2image")
            logging.warning("Download Poppler from: https://github.com/oschwartz10612/poppler-windows/releases/")
            return ""
        
        text = ""
        total_pages = 0
        pages_with_text = 0
        try:
            # Convert PDF pages to images, using poppler_path if specified
            if self.poppler_path:
                images = convert_from_path(file_path, poppler_path=self.poppler_path)
            else:
                images = convert_from_path(file_path)
                
            total_pages = len(images)
            for i, image in enumerate(images):
                logging.info(f"Performing OCR on page {i+1} of PDF {file_path}")
                page_text = pytesseract.image_to_string(image)
                if page_text and page_text.strip():
                    text += page_text + "\n"
                    pages_with_text += 1
                    logging.debug(f"Page {i+1} extracted {len(page_text.strip())} characters")
                else:
                    logging.debug(f"Page {i+1} extracted no readable text")
            
            if text.strip():
                logging.info(f"OCR completed: extracted text from {pages_with_text}/{total_pages} pages")
            else:
                logging.info(f"OCR completed: no readable text found in any of the {total_pages} pages")
            
            return text
        except Exception as e:
            # Check if the error is related to Poppler not being available
            error_msg = str(e).lower()
            if any(keyword in error_msg for keyword in ['poppler', 'unable to get page count', 'pdftoppm']):
                logging.warning(f"Poppler is not installed or not in PATH. Cannot perform OCR on PDF {file_path}")
                logging.warning("To enable PDF OCR, please install Poppler:")
                logging.warning("1. Download from: https://github.com/oschwartz10612/poppler-windows/releases/")
                logging.warning("2. Extract and add to PATH, or set POPPLER_PATH environment variable")
                logging.warning("3. Alternatively, use: winget install poppler")
                return ""
            else:
                logging.error(f"Error performing OCR on PDF {file_path}: {str(e)}")
                return ""

    def _has_real_words(self, text, min_real_word_ratio=0.5, min_absolute_real_words=15, min_word_len=3):
        """Check if the text contains a sufficient number of 'real' words using NLTK."""
        if not text or not isinstance(text, str) or not text.strip():
            logging.debug("Text is empty or not a string.")
            return False

        try:
            # Tokenize the text into words
            tokenized_words = word_tokenize(text.lower())
        except Exception as e:
            logging.error(f"Error tokenizing text with NLTK: {e}")
            # Fallback to simple regex if tokenization fails
            words = re.findall(rf'\b[a-zA-Z]{{{min_word_len},}}\b', text.lower())
            real_word_count = len(words)
            logging.debug(f"NLTK tokenization failed. Fallback regex found {real_word_count} potential words.")
            return real_word_count >= min_absolute_real_words

        # Filter out punctuation and very short words, keep only alphabetic tokens
        potential_words = [word for word in tokenized_words if word.isalpha() and len(word) >= min_word_len]

        if not potential_words:
            logging.debug("No potential words found after tokenization and filtering.")
            return False

        # Get the set of English words from NLTK
        english_vocab = set(nltk_words.words())

        # Count how many potential words are actual English words
        real_word_count = 0
        actual_real_words_found = []
        for word in potential_words:
            if word in english_vocab:
                real_word_count += 1
                actual_real_words_found.append(word)
        
        total_potential_words = len(potential_words)
        
        # Calculate the ratio of real words to total potential words
        if total_potential_words == 0:
            current_ratio = 0
        else:
            current_ratio = real_word_count / total_potential_words

        logging.debug(f"NLTK check: Found {real_word_count} real words (e.g., {actual_real_words_found[:10]}) out of {total_potential_words} potential words. Ratio: {current_ratio:.2f}")

        # Check against both absolute minimum and minimum ratio
        if real_word_count >= min_absolute_real_words and current_ratio >= min_real_word_ratio:
            logging.debug(f"Text considered to have enough real words. (Real: {real_word_count} >= {min_absolute_real_words}, Ratio: {current_ratio:.2f} >= {min_real_word_ratio})")
            return True
        elif real_word_count >= min_absolute_real_words and total_potential_words <= 10:  # If few words, but they are real
            logging.debug(f"Text considered to have enough real words based on absolute count for short text. (Real: {real_word_count} >= {min_absolute_real_words})")
            return True
        else:
            logging.debug(f"Text considered NOT to have enough real words. (Real: {real_word_count} < {min_absolute_real_words} or Ratio: {current_ratio:.2f} < {min_real_word_ratio})")
            return False
    
    def _extract_text_from_docx(self, file_path):
        """Extract text from DOCX file."""
        text = ""
        try:
            doc = Document(file_path)
            # Extract paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                text += "\n"
            return text
        except Exception as e:
            logging.error(f"Error processing DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_doc(self, file_path):
        """Extract text from DOC using external tools only (no Microsoft Word COM).

        Priority order:
        1) antiword (fast, simple)
        2) LibreOffice headless convert-to TXT
        3) LibreOffice headless convert-to DOCX then parse with python-docx
        """
        logging.info(f"Processing DOC using external tools (antiword/LibreOffice): {file_path}")

        # 1) antiword primary
        text = self._doc_with_antiword(file_path)
        if text and text.strip():
            return text

        # 2) LibreOffice -> TXT
        text = self._doc_to_txt_with_libreoffice(file_path)
        if text and text.strip():
            return text

        # 3) LibreOffice -> DOCX -> parse
        text = self._doc_to_docx_with_libreoffice_then_parse(file_path)
        if text and text.strip():
            return text

        logging.warning("All .doc extraction methods failed (antiword/LibreOffice). Returning empty text.")
        return ""

    def _extract_doc_with_fallbacks(self, file_path):
        """Deprecated: legacy fallback chain retained for compatibility.

        New primary flow is in _extract_text_from_doc (antiword -> LibreOffice TXT -> LibreOffice DOCX).
        """
        return self._extract_text_from_doc(file_path)

    def _doc_to_txt_with_libreoffice(self, file_path):
        """Use LibreOffice headless conversion to produce a plain text file."""
        soffice = self.soffice_path or 'soffice'
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                cmd = [
                    soffice, '--headless', '--convert-to', 'txt:Text', '--outdir', tmpdir, file_path
                ]
                logging.info(f"Attempting LibreOffice TXT conversion for .doc: {file_path}")
                proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
                if proc.returncode != 0:
                    logging.debug(f"LibreOffice TXT conversion stderr: {proc.stderr.strip()}")
                    return ""
                # Find the output .txt file (same base name)
                base = os.path.splitext(os.path.basename(file_path))[0]
                out_path = os.path.join(tmpdir, base + '.txt')
                if os.path.isfile(out_path):
                    with open(out_path, 'r', encoding='utf-8', errors='ignore') as f:
                        data = f.read()
                    logging.info(f"LibreOffice TXT conversion succeeded for {file_path}, {len(data.strip())} chars")
                    return data
                return ""
        except FileNotFoundError:
            # soffice not found
            logging.debug("LibreOffice 'soffice' not found. Skipping TXT conversion.")
            return ""
        except subprocess.TimeoutExpired:
            logging.warning("LibreOffice TXT conversion timed out.")
            return ""
        except Exception as e:
            logging.debug(f"LibreOffice TXT conversion error: {e}")
            return ""

    def _doc_to_docx_with_libreoffice_then_parse(self, file_path):
        """Convert .doc to .docx via LibreOffice then parse with python-docx."""
        soffice = self.soffice_path or 'soffice'
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                cmd = [
                    soffice, '--headless', '--convert-to', 'docx:Office Open XML Text', '--outdir', tmpdir, file_path
                ]
                logging.info(f"Attempting LibreOffice DOCX conversion for .doc: {file_path}")
                proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
                if proc.returncode != 0:
                    logging.debug(f"LibreOffice DOCX conversion stderr: {proc.stderr.strip()}")
                    return ""
                base = os.path.splitext(os.path.basename(file_path))[0]
                out_path = os.path.join(tmpdir, base + '.docx')
                if os.path.isfile(out_path):
                    try:
                        return self._extract_text_from_docx(out_path)
                    except Exception as e:
                        logging.debug(f"Parsing converted DOCX failed: {e}")
                        return ""
                return ""
        except FileNotFoundError:
            logging.debug("LibreOffice 'soffice' not found. Skipping DOCX conversion.")
            return ""
        except subprocess.TimeoutExpired:
            logging.warning("LibreOffice DOCX conversion timed out.")
            return ""
        except Exception as e:
            logging.debug(f"LibreOffice DOCX conversion error: {e}")
            return ""

    def _doc_with_antiword(self, file_path):
        """Use 'antiword' CLI if available to extract text from .doc."""
        try:
            # Check availability
            avail = subprocess.run(['antiword', '-h'], capture_output=True, text=True)
            if avail.returncode not in (0, 1):  # some builds return 1 on -h
                return ""
        except FileNotFoundError:
            return ""
        except Exception:
            return ""

        try:
            proc = subprocess.run(['antiword', '-w', '0', file_path], capture_output=True)
            if proc.returncode == 0 and proc.stdout:
                try:
                    text = proc.stdout.decode('utf-8', errors='ignore') if isinstance(proc.stdout, (bytes, bytearray)) else proc.stdout
                except Exception:
                    text = proc.stdout if isinstance(proc.stdout, str) else proc.stdout.decode('latin-1', errors='ignore')
                logging.info(f"antiword extracted {len(text.strip())} chars from {file_path}")
                return text
            else:
                logging.debug(f"antiword failed for {file_path} with code {proc.returncode}")
                return ""
        except Exception as e:
            logging.debug(f"antiword error: {e}")
            return ""
    
    def _extract_text_from_xlsx(self, file_path):
        """Extract text from XLSX file."""
        try:
            df = pd.read_excel(file_path, header=None, dtype=str)
            df = df.fillna('')
            text = ''
            for row in df.itertuples(index=False):
                text += ' '.join(str(cell) for cell in row) + "\n"
            return text
        except Exception as e:
            logging.error(f"Error processing XLSX {file_path}: {str(e)}")
            return ""
    
    def _analyze_pdf_with_vision_llm(self, file_path):
        """Analyze PDF with vision-capable LLM when OCR fails."""
        if not self.pdf_ocr_available:
            logging.warning("Cannot analyze PDF with vision LLM: pdf2image not available.")
            return ""
        
        if not self.llm_analyzer or not self.llm_analyzer.vision_llm:
            logging.warning("Vision LLM not available in LLM analyzer.")
            return ""
        
        try:
            # Convert PDF pages to images
            if self.poppler_path:
                images = convert_from_path(file_path, poppler_path=self.poppler_path)
            else:
                images = convert_from_path(file_path)
            
            if not images:
                logging.warning(f"Could not convert PDF {file_path} to images for vision analysis.")
                return ""
            
            all_text = ""
            
            # Analyze each page with the vision model
            for i, image in enumerate(images[:5]):  # Limit to first 5 pages to avoid excessive processing
                try:
                    # Save image temporarily for LLM analysis
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                        image.save(temp_file.name, 'PNG')
                        temp_image_path = temp_file.name
                    
                    # Create prompt for vision analysis
                    prompt = f"""Analyze this image from page {i+1} of a PDF document. Please:

1. Extract ALL visible text, including any text that might be handwritten, stylized, or difficult for OCR to read
2. Describe the type of document this appears to be
3. Identify any key information like names, dates, amounts, or important details
4. If no readable text is visible, provide a detailed description of what you can see

Please be thorough in extracting text - include headers, body text, captions, and any other readable content."""

                    # Use the configured vision model to analyze the image with robust fallback
                    response_text = self.llm_analyzer.analyze_image_text(temp_image_path, prompt)
                    
                    if response_text and response_text.strip():
                        all_text += f"[Page {i+1} Vision Analysis]\n{response_text.strip()}\n\n"
                        logging.info(f"Vision LLM extracted text from page {i+1}: {len(response_text.strip())} characters")
                    
                    # Clean up temporary file
                    try:
                        os.unlink(temp_image_path)
                    except:
                        pass
                        
                except Exception as e:
                    logging.error(f"Error analyzing page {i+1} with vision LLM: {e}")
                    continue
            
            if all_text.strip():
                logging.info(f"Vision LLM analysis completed for {file_path}: {len(all_text.strip())} total characters extracted")
                return all_text.strip()
            else:
                logging.warning(f"Vision LLM analysis completed but no text extracted from {file_path}")
                return ""
                
        except Exception as e:
            logging.error(f"Error in vision LLM analysis for {file_path}: {e}")
            return ""
    
    def _analyze_image_with_vision_llm(self, file_path):
        """Analyze image with vision-capable LLM when OCR fails."""
        if not self.llm_analyzer or not self.llm_analyzer.vision_llm:
            logging.warning("Vision LLM not available in LLM analyzer.")
            return ""
        
        try:
            # Create prompt for vision analysis
            prompt = """Analyze this image carefully. Please:

1. Extract ALL visible text, including any text that might be handwritten, stylized, watermarked, or difficult for OCR to read
2. Describe the type of document or image this appears to be
3. Identify any key information like names, dates, amounts, addresses, or important details
4. If no readable text is visible, provide a detailed description of what you can see

Please be thorough in extracting text - include any headers, body text, captions, labels, and any other readable content, no matter how small or stylized."""

            # Use the configured vision model via analyzer helper (with fallback)
            response_text = self.llm_analyzer.analyze_image_text(str(file_path), prompt)
            
            if response_text and response_text.strip():
                logging.info(f"Vision LLM extracted text from image {file_path}: {len(response_text.strip())} characters")
                return f"[Vision LLM Analysis]\n{response_text.strip()}"
            else:
                logging.warning(f"Vision LLM analysis completed but no text extracted from {file_path}")
                return ""
                
        except Exception as e:
            logging.error(f"Error in vision LLM analysis for image {file_path}: {e}")
            return ""
